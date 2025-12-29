import os
from typing import List, Dict, Tuple, Optional
import re
from pathlib import Path

# PDF processing
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    from pdfminer.high_level import extract_text as pdfminer_extract_text
except ImportError:
    pdfminer_extract_text = None

# DOCX processing
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# TXT processing
try:
    import chardet
except ImportError:
    chardet = None


class DocumentProcessor:
    """Processes various document formats and prepares them for RAG"""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize document processor
        
        Args:
            chunk_size: Target size for each text chunk (in characters)
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.supported_formats = ['.pdf', '.docx', '.txt', '.md']
    
    def is_supported_format(self, file_path: str) -> bool:
        """Check if file format is supported"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_formats
    
    def extract_text_from_pdf(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Tuple of (full_text, page_metadata)
        """
        if not PyPDF2:
            raise ImportError("PyPDF2 is not installed. Install it with: pip install PyPDF2")
        
        text_parts = []
        page_metadata = []
        
        try:
            # Try PyPDF2 first (faster but sometimes less accurate)
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                num_pages = len(pdf_reader.pages)
                
                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    
                    if text.strip():
                        text_parts.append(text)
                        page_metadata.append({
                            'page': page_num + 1,
                            'total_pages': num_pages,
                            'char_start': len(''.join(text_parts[:-1])),
                            'char_end': len(''.join(text_parts))
                        })
        except Exception as e:
            if pdfminer_extract_text:
                print(f"PyPDF2 failed, trying pdfminer: {e}")
                # Fallback to pdfminer if PyPDF2 fails
                try:
                    text = pdfminer_extract_text(file_path)
                    text_parts = [text]
                    page_metadata = [{'page': 1, 'total_pages': 1, 'char_start': 0, 'char_end': len(text)}]
                except Exception as e2:
                    raise Exception(f"Failed to extract PDF text: {e2}")
            else:
                raise Exception(f"Failed to extract PDF text and pdfminer is not installed: {e}")
        
        full_text = '\n\n'.join(text_parts)
        return full_text, page_metadata
    
    def extract_text_from_docx(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Tuple of (full_text, paragraph_metadata)
        """
        if not DocxDocument:
            raise ImportError("python-docx is not installed. Install it with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            para_metadata = []
            
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    text_parts.append(text)
                    para_metadata.append({
                        'paragraph': i + 1,
                        'char_start': len('\n\n'.join(text_parts[:-1])),
                        'char_end': len('\n\n'.join(text_parts))
                    })
            
            full_text = '\n\n'.join(text_parts)
            return full_text, para_metadata
            
        except Exception as e:
            raise Exception(f"Failed to extract DOCX text: {e}")
    
    def extract_text_from_txt(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            Tuple of (full_text, metadata)
        """
        try:
            # Detect encoding if chardet is available
            encoding = 'utf-8'
            if chardet:
                with open(file_path, 'rb') as file:
                    raw_data = file.read()
                    result = chardet.detect(raw_data)
                    encoding = result['encoding'] or 'utf-8'
            
            # Read with detected encoding
            with open(file_path, 'r', encoding=encoding, errors='ignore') as file:
                text = file.read()
            
            metadata = [{'format': 'txt', 'char_start': 0, 'char_end': len(text)}]
            return text, metadata
            
        except Exception as e:
            raise Exception(f"Failed to extract TXT text: {e}")
    
    def extract_text(self, file_path: str) -> Tuple[str, List[Dict]]:
        """
        Extract text from any supported document format
        
        Args:
            file_path: Path to document file
            
        Returns:
            Tuple of (full_text, metadata)
        """
        ext = Path(file_path).suffix.lower()
        
        if ext == '.pdf':
            return self.extract_text_from_pdf(file_path)
        elif ext == '.docx':
            return self.extract_text_from_docx(file_path)
        elif ext in ['.txt', '.md']:
            return self.extract_text_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def clean_text(self, text: str) -> str:
        """
        Clean and normalize extracted text
        
        Args:
            text: Raw text to clean
            
        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove excessive newlines (keep paragraph breaks)
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Remove special characters that might cause issues
        text = text.replace('\x00', '')
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def split_into_sentences(self, text: str) -> List[str]:
        """
        Split text into sentences
        
        Args:
            text: Text to split
            
        Returns:
            List of sentences
        """
        # Simple sentence splitting 
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def create_chunks(self, text: str, metadata: Optional[List[Dict]] = None) -> List[Dict]:
        """
        Split text into overlapping chunks
        
        Args:
            text: Text to chunk
            metadata: Optional source metadata
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        # Clean text first
        text = self.clean_text(text)
        
        chunks = []
        start = 0
        text_length = len(text)
        chunk_id = 0
        
        while start < text_length:
            # Calculate end position
            end = start + self.chunk_size
            
            # If we're not at the end, try to break at a sentence boundary
            if end < text_length:
                # Look for sentence ending punctuation
                last_period = text.rfind('.', start, end)
                last_exclamation = text.rfind('!', start, end)
                last_question = text.rfind('?', start, end)
                
                # Find the last sentence boundary
                break_point = max(last_period, last_exclamation, last_question)
                
                # If we found a good break point, use it
                if break_point > start + (self.chunk_size // 2):  # Ensure chunk isn't too small
                    end = break_point + 1
            
            # Extract chunk
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                chunk_data = {
                    'text': chunk_text,
                    'chunk_id': chunk_id,
                    'char_start': start,
                    'char_end': end,
                    'chunk_size': len(chunk_text)
                }
                
                # Add source metadata if provided
                if metadata:
                    # Find which metadata entry this chunk belongs to
                    for meta in metadata:
                        if 'char_start' in meta and 'char_end' in meta:
                            if start >= meta['char_start'] and start < meta['char_end']:
                                chunk_data.update({k: v for k, v in meta.items() 
                                                 if k not in ['char_start', 'char_end']})
                                break
                
                chunks.append(chunk_data)
                chunk_id += 1
            
            # Move start position (with overlap)
            start = end - self.chunk_overlap
            
            # Prevent infinite loop
            if start <= 0 or end >= text_length:
                break
        
        return chunks
    
    def process_document(self, file_path: str) -> Tuple[List[str], List[Dict], Dict]:
        """
        Complete document processing pipeline
        
        Args:
            file_path: Path to document
            
        Returns:
            Tuple of (chunk_texts, chunk_metadata, document_stats)
        """
        # Extract text
        full_text, source_metadata = self.extract_text(file_path)
        
        # Create chunks
        chunks = self.create_chunks(full_text, source_metadata)
        
        # Separate text and metadata
        chunk_texts = [chunk['text'] for chunk in chunks]
        chunk_metadata = [{k: v for k, v in chunk.items() if k != 'text'} 
                         for chunk in chunks]
        
        # Calculate statistics
        stats = {
            'total_characters': len(full_text),
            'total_chunks': len(chunks),
            'avg_chunk_size': sum(len(c) for c in chunk_texts) / len(chunk_texts) if chunks else 0,
            'file_format': Path(file_path).suffix.lower(),
            'source_segments': len(source_metadata)
        }
        
        return chunk_texts, chunk_metadata, stats
    
    def extract_keywords(self, text: str, max_keywords: int = 10) -> List[str]:
        """
        Extract potential keywords from text (simple frequency-based approach)
        
        Args:
            text: Text to extract keywords from
            max_keywords: Maximum number of keywords to return
            
        Returns:
            List of keywords
        """
        # Clean and tokenize
        text = self.clean_text(text.lower())
        words = re.findall(r'\b[a-z]{4,}\b', text)  # Words with 4+ characters
        
        # Common stop words to filter out
        stop_words = {
            'this', 'that', 'these', 'those', 'with', 'from', 'have', 'been',
            'were', 'will', 'would', 'could', 'should', 'their', 'there', 'where',
            'which', 'what', 'when', 'about', 'also', 'into', 'through', 'during',
            'before', 'after', 'above', 'below', 'between', 'under', 'again', 'further',
            'then', 'once', 'here', 'more', 'most', 'other', 'some', 'such', 'only',
            'same', 'than', 'very', 'just', 'your', 'they', 'them', 'their'
        }
        
        # Count frequencies
        word_freq = {}
        for word in words:
            if word not in stop_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Sort by frequency and return top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [word for word, freq in sorted_words[:max_keywords]]
        
        return keywords
    
    def get_document_summary_stats(self, file_path: str) -> Dict:
        """
        Get summary statistics about a document without full processing
        
        Args:
            file_path: Path to document
            
        Returns:
            Dictionary of statistics
        """
        try:
            file_size = os.path.getsize(file_path)
            file_format = Path(file_path).suffix.lower()
            
            stats = {
                'file_size': file_size,
                'file_format': file_format,
                'file_name': Path(file_path).name
            }
            
            # Add format-specific stats
            if file_format == '.pdf' and PyPDF2:
                try:
                    with open(file_path, 'rb') as file:
                        pdf_reader = PyPDF2.PdfReader(file)
                        stats['num_pages'] = len(pdf_reader.pages)
                except:
                    pass
            elif file_format == '.docx' and DocxDocument:
                try:
                    doc = DocxDocument(file_path)
                    stats['num_paragraphs'] = len(doc.paragraphs)
                except:
                    pass
            
            return stats
            
        except Exception as e:
            return {'error': str(e)}


def process_uploaded_file(uploaded_file, upload_dir: str = "uploads") -> Tuple[str, str]:
    """
    Save an uploaded Streamlit file to disk
    
    Args:
        uploaded_file: Streamlit UploadedFile object
        upload_dir: Directory to save uploads
        
    Returns:
        Tuple of (file_path, file_type)
    """
    # Create upload directory if it doesn't exist
    os.makedirs(upload_dir, exist_ok=True)
    
    # Generate unique filename
    import uuid
    file_extension = Path(uploaded_file.name).suffix
    unique_filename = f"{uuid.uuid4()}{file_extension}"
    file_path = os.path.join(upload_dir, unique_filename)
    
    # Save file
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())
    
    return file_path, file_extension


def cleanup_file(file_path: str):
    """Delete a file from disk"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
    except Exception as e:
        print(f"Error deleting file {file_path}: {e}")